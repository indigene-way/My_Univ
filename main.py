#! /usr/bin/python
# -*- coding: utf-8 -*-

import sys, os
import re 
import uuid
import time
import math
import random
import datetime
from datetime import datetime

import sqlite3
from PyQt5.QtSql import *

from PyQt5 import QtCore, QtGui, uic
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *

import numpy as np

import docx
from docx import Document
from docx.shared import Pt

#============================================================================================================DATA BASE
conn = sqlite3.connect('db.db')
# os.system("icacls db.db /grant *S-1-1-0:(D,WDAC)")	
query = conn.cursor()

#=========================================================================== 
# query.execute("DROP TABLE Univ")
try:
    query.execute("SELECT id FROM Univ ORDER BY id DESC")

except:
    conn.execute("""CREATE TABLE Univ (
                    id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE,
                    univ_name VARCHAR(256),
                    univ_formation VARCHAR(1024),
                    univ_capacity VARCHAR(8),
                    univ_tarif VARCHAR(10),
                    univ_recuter VARCHAR(1024),
                    univ_M2 VARCHAR(1024),
                    univ_link VARCHAR(1024),
                    univ_moduls VARCHAR(1024),
                    univ_prerequis VARCHAR(1024),
                    univ_docs VARCHAR(1024),
                    univ_valid VARCHAR(10),
                    univ_select VARCHAR(4),
                    univ_competance VARCHAR(1024),
                    univ_motiv VARCHAR(256))""")
    
#=========================================================================== 
# query.execute("DROP TABLE Universities")
try:
    query.execute("SELECT id FROM Universities ORDER BY id DESC")

except:
    conn.execute("""CREATE TABLE Universities (
                    id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE,
                    univ_name VARCHAR(256),
                    univ_formation VARCHAR(1024),
                    univ_capacity VARCHAR(8),
                    univ_tarif VARCHAR(10),
                    univ_recuter VARCHAR(1024),
                    univ_M2 VARCHAR(1024),
                    univ_link VARCHAR(1024),
                    univ_moduls VARCHAR(1024),
                    univ_prerequis VARCHAR(1024),
                    univ_docs VARCHAR(1024))""")
                    
#=========================================================================== 
# query.execute("DROP TABLE SelectedU")
try:
    query.execute("SELECT id FROM SelectedU ORDER BY id DESC")

except:
    conn.execute("""CREATE TABLE SelectedU (
                    id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE,
                    univ_name VARCHAR(256),
                    univ_formation VARCHAR(1024),
                    univ_contact VARCHAR(1024),
                    univ_valid VARCHAR(10),
                    univ_select VARCHAR(4),
                    univ_tarif_cap VARCHAR(1024),
                    univ_competance VARCHAR(1024),
                    univ_motiv VARCHAR(256))""")

#======================================================CLASSES
#============================================================================================================MESSAGE FACTORY
class MessageFactory():
    		
	def raiseAdder(self,data):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Information)
		msg.setWindowTitle("OPERATION REUSSIE !")
		msg.setText(data+" ajouté(e) avec succés !")
		msg.exec_()	
		
	def raiseModifier(self,data):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Information)
		msg.setWindowTitle("OPERATION REUSSIE !")
		msg.setText(data+" modifié(e) avec succés !")
		msg.exec_()	
		
	def raiseDeleter(self,data):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Information)
		msg.setWindowTitle("OPERATION REUSSIE !")
		msg.setText(data+" suprimé(e) avec succés !")
		msg.exec_()	
		
	def raiseCaseExcept(self):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Warning)
		msg.setWindowTitle("CASE OBLIGATOIRE VIDE !")
		msg.setText("Veuillez saisir toutes les cases ")
		msg.exec_()
		
	def raiseUSelecter(self):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Warning)
		msg.setWindowTitle("CASE OBLIGATOIRE VIDE !")
		msg.setText("Veuillez selectionner l'université ")
		msg.exec_()

#=========================================================================== 

# =============UI CLASS Widget 		
qtAN= "DESIGN/Univ.ui"
Ui_Univ, QtBaseClass = uic.loadUiType(qtAN)	
class Univ(QDialog, Ui_Univ):#EDIT : MODIF Product Name,Price DIALOG

    def __init__(self):
        QDialog.__init__(self)
        Ui_Univ.__init__(self)
        self.setupUi(self)
        
        self.Fill()
        
    def Fill(self):
        # pass
        widgetl = [self.widget_1,self.widget_2,self.widget_3,self.widget_4,self.widget_5,self.widget_6,self.widget_7]
        j = 0
        while j<=6 :
            widgetl[j].setVisible(False)
            j+=1
        
        label = [self.label_1,self.label_2,self.label_3,self.label_4,self.label_5,self.label_6,self.label_7]
        fl = [self.fl_1,self.fl_2,self.fl_3,self.fl_4,self.fl_5,self.fl_6,self.fl_7]
        etatl = [self.etatl_1,self.etatl_2,self.etatl_3,self.etatl_4,self.etatl_5,self.etatl_6,self.etatl_7]#(valid + Select) 
        tarifl = [self.tarifl_1,self.tarifl_2,self.tarifl_3,self.tarifl_4,self.tarifl_5,self.tarifl_6,self.tarifl_7]
        capl =  [self.capl_1,self.capl_2,self.capl_3,self.capl_4,self.capl_5,self.capl_6,self.capl_7] 
        scrollArea =  [self.scrollArea_1,self.scrollArea_2,self.scrollArea_3,self.scrollArea_4,self.scrollArea_5,self.scrollArea_6,self.scrollArea_7]#(contact)
        debouchel = [self.debouchel_1,self.debouchel_2,self.debouchel_3,self.debouchel_4,self.debouchel_5,self.debouchel_6,self.debouchel_7]
        linkl = [self.linkl_1,self.linkl_2,self.linkl_3,self.linkl_4,self.linkl_5,self.linkl_6,self.linkl_7]
        fondaScroll = [self.fondaScroll_1,self.fondaScroll_2,self.fondaScroll_3,self.fondaScroll_4,self.fondaScroll_5,self.fondaScroll_6,self.fondaScroll_7]
        preqScroll = [self.preqScroll_1,self.preqScroll_2,self.preqScroll_3,self.preqScroll_4,self.preqScroll_5,self.preqScroll_6,self.preqScroll_7]
        competances = [self.comp_1,self.comp_2,self.comp_3,self.comp_4,self.comp_5,self.comp_6,self.comp_7]
        motiv = [self.motiv_1,self.motiv_2,self.motiv_3,self.motiv_4,self.motiv_5,self.motiv_6,self.motiv_7]
        doc = [self.doc_1,self.doc_2,self.doc_3,self.doc_4,self.doc_5,self.doc_6,self.doc_7]
        
        query.execute("SELECT univ_name, univ_formation, univ_valid, univ_select, univ_tarif, univ_capacity, univ_recuter, univ_M2, univ_link, univ_moduls,\
        univ_prerequis, univ_competance, univ_motiv, univ_docs FROM Univ WHERE univ_valid != '...'  ORDER BY univ_valid ASC")
        univ = query.fetchall()
        
        i = 0
        for u in univ :
            query.execute("UPDATE Univ SET univ_valid = (\"{0}\") WHERE univ_name = \"{1}\" "\
                .format("...",str("Université de Strasbourg - Faculté des lettres")))
            
            widgetl[i].setVisible(True)
            
            label[i].setText(str(univ[i][0]).strip("','"))
            fl[i].setText(str(univ[i][1]).strip("','"))
            etatl[i].setText("N° : "+str(univ[i][2]).strip("','")+"\tSelectivité : "+str(univ[i][3]).strip("','"))
            tarifl[i].setText("Tarifs : "+str(univ[i][4]).strip("','"))
            capl[i].setText("Capacité : "+str(univ[i][5]).strip("','"))
            scrollArea[i].setText(str(univ[i][6]).strip("','"))
            debouchel[i].setText(str(univ[i][7]).strip("','"))
            linkl[i].setText(str(univ[i][8]).strip("','"))
            fondaScroll[i].setText(str(univ[i][9]).strip("','"))
            preqScroll[i].setText(str(univ[i][10]).strip("','"))
            competances[i].setText(str(univ[i][11]).strip("','"))
            motiv[i].setText(str(univ[i][12]).strip("','"))
            doc[i].setText(str(univ[i][13]).strip("','"))
            i+=1
        
#=========================================================================== 

# =============UI CLASS Widget 		
qtANPR= "DESIGN/Choix.ui"
Ui_Choix, QtBaseClass = uic.loadUiType(qtANPR)	
class Choix(QDialog, Ui_Choix):#EDIT : MODIF Product Name,Price DIALOG

    def __init__(self):
        QDialog.__init__(self)
        Ui_Choix.__init__(self)
        self.setupUi(self)
        
        self.title.setVisible(False)
        self.titleIcon.setVisible(False)
        
        self.spinBox.setVisible(False)
    
        self.Line1.setVisible(False)
        self.Line2.setVisible(False)
        self.Line3.setVisible(False)
        self.Line4.setVisible(False)
        self.Line5.setVisible(False)
        self.Line6.setVisible(False)
        self.Line7.setVisible(False)
        
        self.lineList = [self.Line1, self.Line2, self.Line3, self.Line4, self.Line5, self.Line6, self.Line7]
        self.modulList = []
        self.preList = []
        self.docList = []
        
        self.modul = ""
        self.pre = ""
        self.doc = ""

        self.C = [self.c1, self.c2, self.c3, self.c4, self.c5, self.c6, self.c7, self.c8, self.c9, self.c10, self.c11, self.c12, self.c13, self.c14]
                           
        #INSTANCES
        self.fac = Univ()
        self.TableWidgetInit() 
        self.MF = MessageFactory()  
        #SIGNALS MODUL,PRE,DOCS
        self.c12.clicked.connect(self.AddCriteria)
        self.c13.clicked.connect(self.AddCriteria)
        self.c14.clicked.connect(self.AddCriteria)        
        #SIGNALS SPIBBOX
        self.spinBox.valueChanged.connect(self.spinValue)  
        
        #SIGNALS CHOICE
        self.editChoice.clicked.connect(self.Edit)
        self.myChoice.clicked.connect(self.SaveChoice)
        
        #SIGNALS 
        self.addInfo.clicked.connect(self.spinValue)
        self.addUniv.clicked.connect(self.AddUniv)
        self.addInfo.clicked.connect(self.AddData)       
        self.newUniv.clicked.connect(self.CleanU)       
        self.deleteUniv.clicked.connect(self.DeleteUniv)       

    def CleanU(self):
        i = 0
        while i <= 10 :
            self.C[i].setText("")
            i+=1
        self.title.setText("Modules, Prérequis, Docs")
        
    def TableWidgetInit(self):
        #CLEAN TABLE WIDGET
        self.tableWidget.clear()
        self.tableWidget.clearContents()
        self.tableWidget.setRowCount(0)
        font = QFont()
        font.setBold(True)

        self.tableWidget.setColumnCount(7)
        self.tableWidget.setColumnWidth(0, 50)
        self.tableWidget.setColumnWidth(1, 250)
        self.tableWidget.setColumnWidth(2, 250)
        self.tableWidget.setColumnWidth(3, 60)
        self.tableWidget.setColumnWidth(4, 50)
        self.tableWidget.setColumnWidth(5, 105)
        self.tableWidget.setColumnWidth(6, 250)

        self.tableWidget.setHorizontalHeaderLabels(['Classée', 'Université', 'Formation','Capacité','Tarifs','Contacte','Master 2'])
        self.header = self.tableWidget.horizontalHeader()
        self.header.setFont(font)
        self.header.setDefaultAlignment(Qt.AlignHCenter)
                      
        query.execute("SELECT univ_valid,univ_name, univ_formation, univ_capacity, univ_tarif, univ_recuter, univ_M2\
        FROM Univ ORDER BY univ_valid ASC")
        for row, form in enumerate(query):
            self.tableWidget.insertRow(row)
            for column, item in enumerate(form):
            #print(str(item))
                self.tableWidget.setItem(row, column,QTableWidgetItem(str(item))) 
        self.fac.Fill()
                         
    def TableWidgetInitChoice(self):
        #CLEAN TABLE WIDGET
        self.tableWidget.clear()
        self.tableWidget.clearContents()
        self.tableWidget.setRowCount(0)
        font = QFont()
        font.setBold(True)

        self.tableWidget.setColumnCount(7)
        self.tableWidget.setColumnWidth(0, 50)
        self.tableWidget.setColumnWidth(1, 250)
        self.tableWidget.setColumnWidth(2, 250)
        self.tableWidget.setColumnWidth(3, 250)
        self.tableWidget.setColumnWidth(4, 50)
        self.tableWidget.setColumnWidth(5, 80)
        self.tableWidget.setColumnWidth(6, 85)

        self.tableWidget.setHorizontalHeaderLabels(['Classée', 'Université', 'Formation','Contacte','Select','Capacité','Compétances'])
        
        self.header = self.tableWidget.horizontalHeader()
        self.header.setFont(font)
        self.header.setDefaultAlignment(Qt.AlignHCenter)    
                      
        query.execute("SELECT DISTINCT univ_valid,univ_name, univ_formation,univ_recuter,univ_select,univ_capacity, univ_competance\
        FROM Univ WHERE univ_valid != '...' ORDER BY univ_valid ASC")
        for row, form in enumerate(query):
            self.tableWidget.insertRow(row)
            for column, item in enumerate(form):
            #print(str(item))
                self.tableWidget.setItem(row, column,QTableWidgetItem(str(item)))  
        self.fac.Fill()             

    def spinValue(self):#SLOT TO SPIBBOX VALUE CHANGES
        i = self.spinBox.value()
        if i > 0 :
        
            if self.p <= i :
                self.addInfo.setVisible(True)
                while self.p <= i : 
                    try :
                        self.lineList[self.p - 1].setVisible(True)
                    except :
                        pass
                    self.p += 1
                    i = self.spinBox.value()

                    
            elif  self.p >= i and i != 0: 
                while self.p > i : 
                    try :
                        self.lineList[self.p - 1].setVisible(False)
                    except :
                        pass
                    self.p -= 1
                    i = self.spinBox.value()
                             
        else :        
            self.Line1.setVisible(False)
            self.Line2.setVisible(False)
            self.Line3.setVisible(False)
            self.Line4.setVisible(False)
            self.Line5.setVisible(False) 
            self.Line6.setVisible(False) 
            
            self.Line1.clear()
            self.Line2.clear()
            self.Line3.clear()
            self.Line4.clear()
            self.Line5.clear()
            self.Line6.clear()
                    
            self.modulList = []
            self.preList = []
            self.docList = []

    def AddCriteria(self):#SET VISIBLE TO EDITLINE WHEN SPIBOX VALUE CHANGES
        self.Line1.clear()
        self.Line2.clear()
        self.Line3.clear()
        self.Line4.clear()
        self.Line5.clear()
        self.Line6.clear()
        self.Line7.clear()
        self.Line1.setVisible(False)
        self.Line2.setVisible(False)
        self.Line3.setVisible(False)
        self.Line4.setVisible(False)
        self.Line5.setVisible(False) 
        self.Line6.setVisible(False) 
        self.Line7.setVisible(False) 
        self.spinBox.setValue(0)
        sender = self.sender()
        
        self.spinBox.setValue(0)
        self.titleIcon.setVisible(True)
        
        if sender == self.c12 :
            self.p = 1
            self.title.setVisible(True)
            self.title.setText("Modules fondamentaux")
            self.spinBox.setVisible(True)
        
        if sender == self.c13 :
            self.p = 1
            self.title.setVisible(True)
            self.title.setText("Prérequis")
            self.spinBox.setVisible(True)
        
        if sender == self.c14 :
            self.p = 1
            self.title.setVisible(True)
            self.title.setText("Dossier requis")
            self.spinBox.setVisible(True)

    def AddData(self):#SLOT TO UPDATE DATA FROM LINES TO DB
        
        #==========HOVER ONN TITLE
        try :
            rowPosition = self.tableWidget.rowCount()
            index = self.tableWidget.currentRow()

            uName = self.tableWidget.item(index,1).text()  
        except :
            self.MF.raiseUSelecter()
    
        if self.title.text() == "Modules fondamentaux" :
            try :
                self.Line1.setVisible(False)
                self.Line2.setVisible(False)
                self.Line3.setVisible(False)
                self.Line4.setVisible(False)
                self.Line5.setVisible(False)
                self.Line6.setVisible(False)
                self.Line7.setVisible(False)
                self.modulList = []
                self.modul = ""
                t = 0
                for i in self.lineList :
                    if i.text() != "" :
                        self.modulList.append(i.text())
                        self.modul += "-"+self.modulList[t] + "\n"
                        t+=1
                query.execute("UPDATE Univ SET univ_moduls = (\"{0}\") WHERE univ_name = \"{1}\" "\
                .format(self.modul,str(uName)))
                conn.commit()
                self.spinBox.setValue(0) 
                self.MF.raiseAdder("Modules") 
            except :
                self.MF.raiseUSelecter()
               
        elif self.title.text() == "Prérequis" :
            try:
                self.Line1.setVisible(False)
                self.Line2.setVisible(False)
                self.Line3.setVisible(False)
                self.Line4.setVisible(False)
                self.Line5.setVisible(False)
                self.Line6.setVisible(False)
                self.Line7.setVisible(False)
                self.preList = []
                self.pre = ""
                t = 0
                for i in self.lineList :
                    if i.text() != "" :
                        self.preList.append(i.text())
                        self.pre += "-"+self.preList[t] + "\n"
                        t+=1
                query.execute("UPDATE Univ SET univ_prerequis = (\"{0}\") WHERE univ_name = \"{1}\" "\
                .format(self.pre,str(uName)))
                conn.commit()
                self.spinBox.setValue(0)  
                self.MF.raiseAdder("Prérequis") 
            except :
                self.MF.raiseUSelecter()  
                
        elif self.title.text() == "Dossier requis" :
            try :
                self.docList = []
                self.doc = ""
                t = 0
                for i in self.lineList :
                    if i.text() != "" :
                        self.docList.append(i.text())
                        self.doc += "-"+self.docList[t] + "\n"
                        t+=1
                query.execute("UPDATE Univ SET univ_docs = (\"{0}\") WHERE univ_name = \"{1}\" "\
                .format(self.doc,str(uName)))
                conn.commit()
                self.spinBox.setValue(0) 
                self.MF.raiseAdder("Dossier")  
            except :
                self.MF.raiseUSelecter()
        else :
            self.MF.raiseUSelecter()
       
    def AddUniv(self):#SLOT TO ADD UNIV INTO DATA
        if self.c1.text() != "" and self.c2.text() != "" and self.c3.text() != "" and self.c4.text() != ""\
        and self.c5.text() != "" and self.c6.text() != "" and self.c7.text() != "":          
            query.execute("INSERT INTO Univ (univ_name, univ_formation, univ_capacity, univ_tarif,\
                univ_recuter, univ_M2, univ_link, univ_valid, univ_select, univ_competance, univ_motiv)\
                VALUES (\"{0}\",\"{1}\",\"{2}\",\"{3}\",\"{4}\",\"{5}\",\"{6}\",\"{7}\",\"{8}\",\"{9}\",\"{10}\")".format(\
                str(self.c1.text()),str(self.c2.text()),str(self.c3.text()),str(self.c4.text()),str(self.c5.text()),str(self.c6.text())\
                ,str(self.c7.text()),str(self.c8.text()),str(self.c9.text()),str(self.c10.text()),str(self.c11.text())))
            conn.commit()
            self.TableWidgetInit()
            self.MF.raiseAdder("Université") 
        else :
                self.MF.raiseCaseExcept()    
          
    def DeleteUniv(self):#SLOT TO ADD UNIV INTO DATA
        try :
            #==========HOVER ONN TITLE
            rowPosition = self.tableWidget.rowCount()
            index = self.tableWidget.currentRow()

            uName = self.tableWidget.item(index,1).text()     
            
            query.execute("DELETE FROM `Univ` WHERE `univ_name` = \"{0}\"".format(str(uName)))               
            conn.commit()
            self.TableWidgetInit()
            self.MF.raiseDeleter("Université") 
        except :
            self.MF.raiseUSelecter() 

#=============================================CHOICE PART

    def Edit(self):#FILLING U DATA SELECTION
        try :
            #==========HOVER ONN TITLE
            rowPosition = self.tableWidget.rowCount()
            index = self.tableWidget.currentRow()
            
            uVld = self.tableWidget.item(index,0).text() 
            uName = self.tableWidget.item(index,1).text() 
            uForm = self.tableWidget.item(index,2).text() 
            uRec = self.tableWidget.item(index,3).text() 
            uSelect = self.tableWidget.item(index,4).text() 
            uCap = self.tableWidget.item(index,5).text() 
            uComp = self.tableWidget.item(index,6).text() 
            
            query.execute("UPDATE Univ SET (univ_valid,univ_formation,univ_recuter,univ_select,univ_capacity,univ_competance) = \
            (\"{0}\",\"{1}\",\"{2}\",\"{3}\",\"{4}\",\"{5}\") \
            WHERE univ_name = \"{6}\" "\
            .format(str(uVld),str(uForm),str(uRec),str(uSelect),str(uCap),str(uComp),str(uName)))
            conn.commit()
            self.TableWidgetInitChoice()
            self.fac.Fill()
            self.MF.raiseModifier("Université") 
        except :
            self.MF.raiseUSelecter()
                   
    def SaveChoice(self):#DOCUMENT DOCX GENERATION
        try :
            rowPosition = self.tableWidget.rowCount()
            index = self.tableWidget.currentRow()
            uName = self.tableWidget.item(index,1).text()  
            uForm = self.tableWidget.item(index,2).text()  
            
            TICKET = Document()
            style = TICKET.styles['Normal']
            font = style.font
            
            h = TICKET.add_heading(uName+"", level=1)
            
            h.bold = True
            h.italic = True
            p = TICKET.add_paragraph(uForm)
            # font.size = Pt(8)        
            
            query.execute("SELECT DISTINCT univ_valid,univ_select,univ_motiv \
            FROM Univ WHERE univ_name = '"+uName+"' ORDER BY id ASC")
            mUniv = query.fetchall()
            p = TICKET.add_heading("Classement : "+str(mUniv[0][0]).strip("(',')")+", Selectivité : "+str(mUniv[0][1]).strip("(',')"), level=3)
                
            query.execute("SELECT DISTINCT univ_capacity,univ_tarif,univ_recuter,univ_M2, univ_link, univ_competance, univ_motiv \
            FROM Univ WHERE univ_name = '"+uName+"' ORDER BY id ASC")
            myUniv = query.fetchall()
            
            p = TICKET.add_paragraph("Débouché/Master : "+str(myUniv[0][3]).strip("(',')"))
                
            p = TICKET.add_paragraph(" ________________________________________________________________________________________________ ")
        
            p = TICKET.add_heading("Informations Générales", level=2)
            tab = TICKET.add_table(1,3)
            heading_cells = tab.rows[0].cells
            heading_cells[0].text = "Capacité d'acceuil"
            heading_cells[1].text = "Tarifs d'niscription"
            heading_cells[2].text = "Chargé de spécialité"
            
            cells = tab.add_row().cells
            cells[0].text = str(myUniv[0][0]).strip("(',')")
            cells[1].text = str(myUniv[0][1]).strip("(',')") + " €"
            cells[2].text = str(myUniv[0][2]).strip("(',')")
            
            p = TICKET.add_heading("Informations Supplémentaires", level=2)        
            p = TICKET.add_paragraph("Compétances : "+str(myUniv[0][5]).strip("(',')"))
            p = TICKET.add_paragraph("\nLien vers la formation : "+str(myUniv[0][4]).strip("(',')"))
            p = TICKET.add_paragraph(" ________________________________________________________________________________________________ ")
        
            p = TICKET.add_heading("Contenu Formation", level=2)
            tab = TICKET.add_table(1,3)
            heading_cells = tab.rows[0].cells
            heading_cells[0].text = "Modules fondamentaux"
            heading_cells[1].text = "Prérequis"
            heading_cells[2].text = "Dossier"
                
            query.execute("SELECT DISTINCT univ_moduls,univ_prerequis,univ_docs \
            FROM Univ WHERE univ_name = '"+uName+"' ORDER BY id ASC")
            myU = query.fetchall()
            
            
            cells = tab.add_row().cells
            cells[0].text = str(myU[0][0]).strip("(',')")
            cells[1].text = str(myU[0][1]).strip("(',')")
            cells[2].text = str(myU[0][2]).strip("(',')")
            p = TICKET.add_paragraph(" ________________________________________________________________________________________________ ")
            
            p = TICKET.add_paragraph("Motivation : "+str(myUniv[0][6]).strip("(',')"))

            TICKET.save("UNIVERSITIES/SELECTED/"+str(mUniv[0][0]).strip("(',')")+"."+uName+".docx" )
            self.MF.raiseAdder("Document") 
        except:
            self.MF.raiseUSelecter()

        
#=========================================================================== 

# =============UI CLASS Widget 		
qtANPRW= "DESIGN/main.ui"
Ui_main, QtBaseClass = uic.loadUiType(qtANPRW)	
class mainWidget(QDialog, Ui_main):#EDIT : MODIF Product Name,Price DIALOG

    def __init__(self):
        QDialog.__init__(self)
        Ui_main.__init__(self)
        self.setupUi(self)
        self.setWindowTitle("Selection Des Universités.")  
          
        
        # Instance:
        self.choice = Choix()
        self.fac = Univ()
        self.Open()                    
        
        # SIGNALS  MENU                
        self.menu1.clicked.connect(self.Menu)
        self.menu2.clicked.connect(self.Menu)
        self.menu3.clicked.connect(self.Menu) 
        
    def Open(self):
        self.dockWidget.setWidget(self.choice)
        self.univ = 0    
        query.execute("SELECT DISTINCT univ_name FROM Univ ORDER BY id ASC")
        myUniv = query.fetchall()
        for u in myUniv :
            self.univ += 1
        self.menu1.setText("Universités ("+str(self.univ)+")")
        
        self.choix = 0    
        query.execute("SELECT DISTINCT univ_name FROM Univ WHERE univ_valid != '...' ORDER BY id ASC")
        mycniv = query.fetchall()
        for u in mycniv :
            self.choix += 1
        self.menu2.setText("Mes Coix ("+str(self.choix)+")")
        
    def Menu(self):
        sender = self.sender()
        
        if sender == self.menu1 :
            self.Open()
            self.choice.TableWidgetInit()  
            
        elif sender == self.menu2 : 
            self.Open()
            self.choice.TableWidgetInitChoice()

        elif sender == self.menu3 :
            self.fac.Fill()
            self.dockWidget.setWidget(self.fac)
     
if __name__ == '__main__':

    app = QApplication(sys.argv)
    ex = mainWidget()        
    ex.show()
    sys.exit(app.exec_())

                    
 
 

